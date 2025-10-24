import io
import docx
import requests
import pandas as pd
import streamlit as st
from app.llm.vectorize_os_annual import run_pipeline
from app.llm.vectorize_os_annual import generate_challenges_report, generate_indicator_tables

st.set_page_config(page_title="AICCRA generator", page_icon="📄")

st.title("📄 AICCRA Report Generator")

tab1, tab2, tab3 = st.tabs(["📊 Annual Report by Indicator", "📋 Indicator Summary Tables", "🎯 Challenges & Lessons Learned"])

with tab1:
    st.header("📊 Annual Report by Indicator")
    st.caption("Generate comprehensive annual reports for specific AICCRA indicators")
    
    indicators = [
        "IPI 1.1",
        "IPI 1.2", 
        "IPI 1.3",
        "IPI 1.4",
        "IPI 2.1",
        "IPI 2.2",
        "IPI 2.3",
        "IPI 3.1",
        "IPI 3.2",
        "IPI 3.3",
        "IPI 3.4",
        "PDO Indicator 1",
        "PDO Indicator 2",
        "PDO Indicator 3",
        "PDO Indicator 4",
        "PDO Indicator 5"
    ]

    selected_indicator = st.selectbox("Select an indicator:", indicators, key="indicator_tab1")
    selected_year = st.selectbox("Select a year:", [2025], key="year_tab1")

    if st.button("Generate Annual Report", type="primary", key="generate_tab1"):
        with st.spinner("Generating annual report..."):
            try:
                url = "https://ia.prms.cgiar.org/api/generate-annual"
                payload = {
                    "indicator": selected_indicator,
                    "year": selected_year,
                    "insert_data": "False"
                }
                response = requests.post(url, json=payload, timeout=600)
                response.raise_for_status()
                report_text = response.json().get("content", "No report found in response.")

                st.session_state.report_text = report_text
                st.session_state.selected_indicator = selected_indicator
                st.session_state.selected_year = selected_year

            except Exception as e:
                st.error(f"⚠️ An error occurred: {e}")

    if "report_text" in st.session_state:
        st.markdown(st.session_state.report_text)
        
        doc = docx.Document()
        for line in st.session_state.report_text.split('\n'):
            doc.add_paragraph(line)
        docx_buffer = io.BytesIO()
        doc.save(docx_buffer)
        docx_buffer.seek(0)

        st.download_button(
            label="📄 Download Report (DOCX)",
            data=docx_buffer,
            file_name=f"annual_report_{st.session_state.selected_indicator}_{st.session_state.selected_year}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="download_tab1_docx"
        )

with tab2:
    st.header("📋 Indicator Summary Tables")
    st.caption("Generate summary tables for all PDO, IPI 1.x, IPI 2.x, and IPI 3.x indicators for 2025.")

    selected_year_tables = st.selectbox("Select a year:", [2025], key="year_tab4")

    if st.button("Generate Tables", type="primary", key="generate_tab4"):
        with st.spinner("Generating summary tables..."):
            try:
                url = "https://ia.prms.cgiar.org/api/generate-annual-tables"
                payload = {
                    "indicator": selected_indicator,
                    "year": selected_year_tables,
                    "insert_data": "False"
                }
                response = requests.post(url, json=payload, timeout=600)
                response.raise_for_status()
                response_data = response.json()
                
                tables = {}
                for group_name, table_list in response_data["tables"].items():
                    tables[group_name] = pd.DataFrame(table_list)
                
                st.session_state.tables = tables
                st.session_state.selected_year_tables = selected_year_tables
            except Exception as e:
                st.error(f"⚠️ An error occurred: {e}")

    if "tables" in st.session_state:
        tables = st.session_state.tables
        for group_name, df_table in tables.items():
            st.subheader(group_name)
            st.dataframe(df_table)
            excel_buffer = io.BytesIO()
            with pd.ExcelWriter(excel_buffer, engine='openpyxl') as writer:
                df_table.to_excel(writer, index=False, sheet_name=group_name)
            excel_buffer.seek(0)
            st.download_button(
                label=f"⬇️ Download {group_name} Table (Excel)",
                data=excel_buffer,
                file_name=f"{group_name}_summary_{st.session_state.selected_year_tables}.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                key=f"download_{group_name}_tab4"
            )

with tab3:
    st.header("🎯 Challenges & Lessons Learned")
    st.caption("Generate cross-cluster reports on challenges faced and lessons learned")
    
    st.info("💡 This report covers all clusters and is not divided by indicators")
    
    selected_year_challenges = st.selectbox("Select a year:", [2025], key="year_tab2")

    if st.button("Generate Challenges Report", type="primary", key="generate_tab2"):
        with st.spinner("Generating challenges and lessons learned report..."):
            try:
                response = generate_challenges_report(selected_year_challenges)
                
                st.session_state.challenges_response = response
                st.session_state.selected_year_challenges = selected_year_challenges
            except Exception as e:
                st.error(f"⚠️ An error occurred: {e}")

    if "challenges_response" in st.session_state:
        st.markdown(st.session_state.challenges_response)

        doc_challenges = docx.Document()
        for line in st.session_state.challenges_response.split('\n'):
            doc_challenges.add_paragraph(line)
        docx_buffer_challenges = io.BytesIO()
        doc_challenges.save(docx_buffer_challenges)
        docx_buffer_challenges.seek(0)

        st.download_button(
            label="📄 Download Challenges Report",
            data=docx_buffer_challenges,
            file_name=f"challenges_lessons_learned_{st.session_state.selected_year_challenges}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="download_tab2_docx"
        )

with st.sidebar:
    st.subheader("📋 Report Types")
    
    st.markdown("""
    **Annual Report by Indicator:**
    - Detailed performance analysis
    - Deliverables and contributions
    - OICRs and innovations
    - Disaggregated targets (selected indicators)
    - Progress tracking and achievements
    
    **Challenges & Lessons Learned:**
    - Cross-cluster analysis
    - Implementation challenges
    - Adaptive strategies
    - Best practices identification
    - Strategic recommendations
    """)
    
    st.divider()
    
    st.subheader("ℹ️ Tips")
    st.markdown("""
    - Annual reports take 2-3 minutes to generate
    - Challenges reports are typically shorter
    - Download buttons appear after generation
    """)